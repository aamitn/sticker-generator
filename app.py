#!/usr/bin/env python3
# ----------------------------------------
# Front & Back Sticker Generator App
# Org : Bitmutex Technologies
# Author : Amit Kumar Nandi
# Version : v0.16
# License : MIT
# **(©) 2025 Bitmutex Technologies. All rights reserved.**
# ----------------------------------------
"""
Professional sticker generator application for UPS and Battery Charger products.
Generates front and back stickers with customizable product details and serial numbers.
"""

import sys
import os
import logging
from typing import Optional, Tuple
from pathlib import Path
from datetime import date

from PyQt6.QtWidgets import (
    QApplication, QWidget, QLabel, QLineEdit, QPushButton,
    QVBoxLayout, QHBoxLayout, QFileDialog, QComboBox, QSpinBox,
    QMessageBox, QGroupBox, QFormLayout, QMainWindow, QMenuBar,
    QCheckBox, QProgressDialog
)
from PyQt6.QtGui import QFont, QIcon, QColor, QPalette, QIntValidator, QAction, QPageSize
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QSettings
from PyQt6.QtPrintSupport import QPrinter, QPrintDialog

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import subprocess
import webbrowser
import platform
import urllib.request
import json
from packaging import version

# ----------------------------------------
# Global Configuration Constants
# ----------------------------------------
class Config:
    """Application configuration constants."""
    # Version
    VERSION = "0.16"
    APP_NAME = "Sticker Generator Tool"
    ORGANIZATION = "Bitmutex"
    GITHUB_REPO = "aamitn/sticker-generator"
    GITHUB_API_URL = f"https://api.github.com/repos/aamitn/sticker-generator/releases/latest"
    
    # Input validation limits
    KVA_MIN = 0
    KVA_MAX = 99999
    UPS_SETS_MIN = 1
    UPS_SETS_MAX = 20
    UPS_PER_SET_MIN = 1
    UPS_PER_SET_MAX = 20
    CHARGERS_MIN = 1
    CHARGERS_MAX = 20
    JOB_OP_MAX = 999999
    
    # File paths - use user directories to avoid permission issues
    APP_DATA_DIR = Path.home() / "Documents" / "Sticker Generator"
    DOCS_DIR = APP_DATA_DIR / "Output"
    LOG_FILE = APP_DATA_DIR / "sticker_generator.log"
    DEFAULT_STICKER = Path.cwd() / "sticker.png"
    
    # Font settings
    HEADING_FONT = "Calibri"
    HEADING_SIZE = 48
    TEXT_FONT = "Calibri"
    BASE_FONT_SIZE = 23
    MIN_FONT_SIZE = 14
    
    # Sticker dimensions
    STICKER_WIDTH = 6.3  # inches


# ----------------------------------------
# Logging Configuration
# ----------------------------------------
def setup_logging():
    """Setup logging configuration with proper error handling."""
    try:
        # Ensure log directory exists
        Config.APP_DATA_DIR.mkdir(parents=True, exist_ok=True)
        
        # Configure logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(Config.LOG_FILE, encoding='utf-8'),
                logging.StreamHandler()
            ]
        )
        logger = logging.getLogger(__name__)
        logger.info(f"Logging initialized. Log file: {Config.LOG_FILE}")
        return logger
    except Exception as e:
        # Fallback to console-only logging if file creation fails
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[logging.StreamHandler()]
        )
        logger = logging.getLogger(__name__)
        logger.warning(f"Could not create log file: {e}. Using console logging only.")
        return logger

logger = setup_logging()


# ----------------------------------------
# Utility Functions
# ----------------------------------------
def fit_text_to_line(
    run, 
    text: str, 
    base_font_size: int = Config.BASE_FONT_SIZE, 
    max_chars_one_line: int = 40, 
    min_font_size: int = Config.MIN_FONT_SIZE
) -> int:
    """
    Dynamically adjust font size to fit text in one line.
    
    Args:
        run: Document run object
        text: Text to fit
        base_font_size: Starting font size
        max_chars_one_line: Maximum characters per line
        min_font_size: Minimum allowed font size
        
    Returns:
        Final font size used
    """
    text_length = len(text)
    font_size = base_font_size
    
    while text_length > max_chars_one_line and font_size > min_font_size:
        font_size -= 1
        max_chars_one_line += 3
    
    run.font.size = Pt(font_size)
    return font_size


def add_page(
    doc: Document, 
    side: str, 
    product_label: str, 
    customer_name: str, 
    serial_number: str, 
    sticker_path: str, 
    show_customer_in_parens: bool = True
) -> None:
    """
    Add a formatted sticker page to the document.
    
    Args:
        doc: Document object
        side: "FRONT SIDE" or "BACK SIDE"
        product_label: Product description
        customer_name: Customer name
        serial_number: Serial number with prefix
        sticker_path: Path to sticker image
        show_customer_in_parens: Whether to show customer name in parentheses
    """
    # Add heading
    heading = doc.add_paragraph()
    run = heading.add_run(side)
    run.font.name = Config.HEADING_FONT
    run.font.size = Pt(Config.HEADING_SIZE)
    run.font.bold = True
    run.font.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")

    # Add sticker image
    if os.path.exists(sticker_path):
        try:
            doc.add_picture(sticker_path, width=Inches(Config.STICKER_WIDTH))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            logger.debug(f"Added sticker image: {sticker_path}")
        except Exception as e:
            logger.error(f"Failed to add sticker image: {e}")
            doc.add_paragraph("[Sticker image error]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        logger.warning(f"Sticker image not found: {sticker_path}")
        doc.add_paragraph("[Sticker image missing]").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Add product name with optional customer name
    if show_customer_in_parens:
        p_name_text = f"{product_label} ({customer_name})"
    else:
        p_name_text = product_label
    
    p_name = doc.add_paragraph(p_name_text)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.runs[0]
    run_name.font.name = Config.TEXT_FONT
    run_name.font.bold = True
    run_name.font.color.rgb = RGBColor(0, 0, 0)
    fit_text_to_line(run_name, p_name_text)

    # Add serial number
    p_serial = doc.add_paragraph(serial_number)
    p_serial.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_serial = p_serial.runs[0]
    run_serial.font.name = Config.TEXT_FONT
    run_serial.font.bold = True
    run_serial.font.color.rgb = RGBColor(0, 0, 0)
    fit_text_to_line(run_serial, serial_number)


def get_financial_year_from_year(year: int) -> str:
    """
    Convert a calendar year to financial year string (YY-YY format).
    
    Args:
        year: Calendar year (e.g., 2025)
        
    Returns:
        Financial year string (e.g., "25-26")
    """
    start = year % 100
    end = (year + 1) % 100
    return f"{start:02d}-{end:02d}"


def get_current_financial_year() -> str:
    """
    Calculate current financial year based on today's date (April-March cycle).
    
    Returns:
        Current financial year string (e.g., "25-26")
    """
    today = date.today()
    year = today.year
    month = today.month
    
    # Financial year starts in April
    if month >= 4:
        fy_start = year
    else:
        fy_start = year - 1
    
    return get_financial_year_from_year(fy_start)


def get_latest_github_release() -> Optional[dict]:
    """
    Fetch the latest release information from GitHub.
    
    Returns:
        Dictionary with 'version' and 'url' or None if failed
    """
    try:
        req = urllib.request.Request(
            Config.GITHUB_API_URL,
            headers={'User-Agent': 'Mozilla/5.0'}
        )
        
        with urllib.request.urlopen(req, timeout=5) as response:
            data = json.loads(response.read().decode())
            
            # Extract version from tag_name (e.g., "v0.15" -> "0.15")
            tag_name = data.get('tag_name', '')
            release_version = tag_name.lstrip('v')
            release_url = data.get('html_url', '')
            
            return {
                'version': release_version,
                'url': release_url
            }
    except Exception as e:
        logger.warning(f"Failed to fetch latest release: {e}")
        return None


# ----------------------------------------
# Worker Thread for DOCX Generation
# ----------------------------------------
class DocxWorker(QThread):
    """
    Background worker thread for generating DOCX files.
    Emits progress updates and handles errors gracefully.
    """
    
    progress = pyqtSignal(int)      # Progress percentage (0-100)
    finished = pyqtSignal(str)      # Final file path
    error = pyqtSignal(str)         # Error message

    def __init__(self, main_window, **kwargs):
        """
        Initialize worker thread.
        
        Args:
            main_window: Reference to main application window
            **kwargs: Generation parameters
        """
        super().__init__()
        self.main_window = main_window
        self.kwargs = kwargs
        logger.info("DocxWorker initialized with parameters")

    def run(self) -> None:
        """Main worker thread execution."""
        try:
            logger.info("Starting DOCX generation")
            
            # Extract parameters
            product_type = self.kwargs.get("product_type", "").upper().strip()
            customer_name = self.kwargs.get("customer_name", "").upper().strip()
            sticker_path = self.kwargs.get("sticker_path", "")
            job_no = self.kwargs.get("job_no", "")
            op_no = self.kwargs.get("op_no", "")
            start_index = self.kwargs.get("start_index", 1)

            # Validate required parameters
            if not all([product_type, customer_name, sticker_path, job_no, op_no]):
                raise ValueError("Missing required parameters")

            doc = Document()
            
            # Determine fiscal year
            if self.main_window.override_fy_cb.isChecked():
                fy_str = self.main_window.fy_dropdown.currentText()
                fy_input_year = 2000 + int(fy_str.split('-')[0])
                fy = get_financial_year_from_year(fy_input_year)
                logger.info(f"Using override FY: {fy}")
            else:
                fy = get_current_financial_year()
                logger.info(f"Using current FY: {fy}")

            # Calculate total pages for progress tracking
            total_pages = self._calculate_total_pages(product_type)
            current_page = 0

            def add_with_progress(*args, **kwargs):
                """Wrapper to track progress."""
                nonlocal current_page
                add_page(*args, **kwargs)
                current_page += 1
                percent = int((current_page / total_pages) * 100) if total_pages > 0 else 0
                self.progress.emit(percent)

            # Generate pages based on product type
            if product_type == "UPS":
                self._generate_ups_stickers(doc, fy, job_no, op_no, customer_name, sticker_path, add_with_progress)
            else:
                self._generate_charger_stickers(doc, fy, job_no, op_no, customer_name, sticker_path, start_index, add_with_progress)

            # Save document
            filename = f"Sticker_{customer_name}_{job_no}_{op_no}_{product_type}.docx"
            output_path = str(self.main_window.save_output_path(filename))
            doc.save(output_path)
            
            logger.info(f"Document saved successfully: {output_path}")
            self.finished.emit(output_path)

        except Exception as e:
            error_msg = f"Generation failed: {str(e)}"
            logger.error(error_msg, exc_info=True)
            self.error.emit(error_msg)

    def _calculate_total_pages(self, product_type: str) -> int:
        """Calculate total pages for progress tracking."""
        if product_type == "UPS":
            num_sets = self.kwargs.get("num_sets", 1)
            ups_per_set = self.kwargs.get("ups_per_set", 1)
            return num_sets * (ups_per_set + 1) * 2  # +1 for BYPASS, 2 sides
        else:
            num_chargers = self.kwargs.get("num_chargers", 1)
            return num_chargers * 2

    def _generate_ups_stickers(self, doc, fy, job_no, op_no, customer_name, sticker_path, add_with_progress):
        """Generate UPS stickers."""
        num_sets = self.kwargs.get("num_sets", 1)
        ups_per_set = self.kwargs.get("ups_per_set", 1)
        kva_rating = self.kwargs.get("kva_rating")
        
        logger.info(f"Generating {num_sets} UPS sets with {ups_per_set} units each")
        
        for set_idx in range(1, num_sets + 1):
            ups_list = [f"UPS{i + 1}" for i in range(ups_per_set)]
            if ups_per_set > 1:
                ups_list.append("BYPASS")

            for unit in ups_list:
                product_label = f"{kva_rating}kVA {unit}"
                serial_number = (
                    f"(SL. NO. : LL/{fy}/{job_no}-OP{op_no}/BYP)" 
                    if unit == "BYPASS" 
                    else f"(SL. NO. : LL/{fy}/{job_no}-OP{op_no}/{unit})"
                )
                
                for side in ["FRONT SIDE", "BACK SIDE"]:
                    add_with_progress(
                        doc, side, product_label, customer_name, 
                        serial_number, sticker_path, True
                    )

    def _generate_charger_stickers(self, doc, fy, job_no, op_no, customer_name, sticker_path, start_index, add_with_progress):
        """Generate Battery Charger stickers."""
        start = 0 if start_index == 0 else 1
        num_chargers = self.kwargs.get("num_chargers", 1)
        voltage = self.kwargs.get("voltage", "")
        current = self.kwargs.get("current", "")
        battery_capacity = self.kwargs.get("battery_capacity", "")
        charger_type = self.kwargs.get("charger_type", "")
        battery_type = self.kwargs.get("battery_type", "")
        
        show_label = self.main_window.show_prod_label_cb.isChecked()
        
        logger.info(f"Generating {num_chargers} charger stickers (show_label={show_label})")

        for i in range(start, num_chargers + start):
            index_label = "" if i == 0 else str(i)

            if show_label:
                product_label = (
                    f"{voltage}V/{current}A {charger_type} "
                    f"for {battery_capacity}Ah {battery_type} battery"
                )
                show_customer_in_parens = True
            else:
                product_label = customer_name
                show_customer_in_parens = False

            serial_number = f"(SL. NO. : LL/{fy}/{job_no}-OP{op_no}/BCH{index_label})"

            for side in ["FRONT SIDE", "BACK SIDE"]:
                add_with_progress(
                    doc, side, product_label, customer_name,
                    serial_number, sticker_path, show_customer_in_parens
                )


# ----------------------------------------
# GUI Main Window
# ----------------------------------------
class StickerApp(QMainWindow):
    """Main application window for sticker generator."""

    def __init__(self):
        """Initialize the application."""
        super().__init__()
        logger.info(f"Initializing {Config.APP_NAME} v{Config.VERSION}")
        
        self.setWindowTitle(Config.APP_NAME)
        self.setWindowIcon(QIcon.fromTheme("document-new"))
        self.setFixedWidth(480)
        self.start_index = 1

        # Initialize UI components
        self.init_ui()
        self.apply_adaptive_theme()
        self.auto_load_sticker()
        self.update_charger_field_state()
        
        # Load persisted settings
        self.settings = QSettings(Config.ORGANIZATION, "StickerGenerator")
        self.load_settings()

    def apply_adaptive_theme(self) -> None:
        """Apply adaptive theme based on system preferences."""
        try:
            app_palette = QApplication.instance().palette()
            base_color = app_palette.color(QPalette.ColorRole.Window)
            brightness = (
                base_color.red() * 0.299 + 
                base_color.green() * 0.587 + 
                base_color.blue() * 0.114
            )
            is_dark = brightness < 128

            palette = QPalette()
            if is_dark:
                palette.setColor(QPalette.ColorRole.Window, QColor("#121212"))
                palette.setColor(QPalette.ColorRole.Base, QColor("#1E1E1E"))
                palette.setColor(QPalette.ColorRole.Text, QColor("#F2F2F2"))
                palette.setColor(QPalette.ColorRole.Button, QColor("#2D89EF"))
                palette.setColor(QPalette.ColorRole.ButtonText, Qt.GlobalColor.white)
            else:
                palette.setColor(QPalette.ColorRole.Window, QColor("#F5F7FB"))
                palette.setColor(QPalette.ColorRole.Base, Qt.GlobalColor.white)
                palette.setColor(QPalette.ColorRole.Text, Qt.GlobalColor.black)
                palette.setColor(QPalette.ColorRole.Button, QColor("#2F80ED"))
                palette.setColor(QPalette.ColorRole.ButtonText, Qt.GlobalColor.white)
            
            QApplication.instance().setPalette(palette)
            self.setPalette(palette)
            self.setFont(QFont("Segoe UI", 10))
            logger.debug(f"Applied {'dark' if is_dark else 'light'} theme")
        
        except Exception as e:
            logger.warning(f"Failed to apply adaptive theme: {e}")

    def auto_load_sticker(self) -> None:
        """Auto-load default sticker image if it exists."""
        # Try multiple locations for the sticker image
        possible_paths = [
            Config.DEFAULT_STICKER,  # Current working directory
            Path(sys.executable).parent / "sticker.png",  # Next to executable (installed)
            Path(sys.argv[0]).parent / "sticker.png",  # Next to script
            Path(__file__).parent / "sticker.png",  # Next to this file
        ]
        
        for sticker_path in possible_paths:
            if sticker_path.exists():
                self.sticker_path.setText(str(sticker_path))
                logger.info(f"Auto-loaded default sticker: {sticker_path}")
                return
        
        logger.info("No default sticker image found")

    def init_ui(self) -> None:
        """Initialize the user interface."""
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout()
        main_layout.setSpacing(15)

        # Create menu bar
        self._create_menu_bar()
        
        # Create UI sections
        fy_box = self._create_fiscal_year_section()
        customer_box = self._create_customer_section()
        sticker_box = self._create_sticker_section()
        ups_box = self._create_ups_section()
        charger_box = self._create_charger_section()
        options_box = self._create_options_section()
        generate_btn = self._create_generate_button()
        
        # Add all sections to main layout
        main_layout.addWidget(fy_box)
        main_layout.addWidget(customer_box)
        main_layout.addWidget(sticker_box)
        main_layout.addWidget(ups_box)
        main_layout.addWidget(charger_box)
        main_layout.addWidget(options_box)
        main_layout.addWidget(generate_btn)
        
        central_widget.setLayout(main_layout)
        
        # Connect signals
        self.product_type.currentTextChanged.connect(self.update_visibility)
        self.update_visibility()

    def _create_menu_bar(self) -> None:
        """Create application menu bar."""
        menu_bar = QMenuBar(self)
        self.setMenuBar(menu_bar)

        # Settings menu
        settings_menu = menu_bar.addMenu("Settings")
        self.start_0_action = QAction("Start BCH numbering from 0", self, checkable=True)
        self.start_0_action.triggered.connect(self.toggle_start_index)
        settings_menu.addAction(self.start_0_action)
        
        self.use_default_printer_action = QAction("Use Default Printer", self, checkable=True)
        settings_menu.addAction(self.use_default_printer_action)

        # Edit menu
        edit_menu = menu_bar.addMenu("Edit")
        open_output_action = QAction("Open Output Path", self)
        open_output_action.triggered.connect(self.open_output_path)
        edit_menu.addAction(open_output_action)
        
        purge_all_action = QAction("Purge All DOCX", self)
        purge_all_action.triggered.connect(self.purge_all_docx)
        edit_menu.addAction(purge_all_action)

        # Help menu
        about_menu = menu_bar.addMenu("Help")
        about_action = QAction("About", self)
        about_action.triggered.connect(self.show_about)
        about_menu.addAction(about_action)
        
        update_action = QAction("Check for Update", self)
        update_action.triggered.connect(self.open_github_release)
        about_menu.addAction(update_action)

    def _create_fiscal_year_section(self) -> QGroupBox:
        """Create fiscal year override section."""
        fy_box = QGroupBox("Fiscal Year")
        fy_layout = QHBoxLayout()
        
        self.override_fy_cb = QCheckBox("Override Fiscal Year")
        self.fy_dropdown = QComboBox()
        
        # Populate FY dropdown ±20 years from current FY
        current_year = date.today().year
        fy_list = [
            f"{y % 100:02d}-{(y + 1) % 100:02d}"
            for y in range(current_year - 20, current_year + 21)
        ]
        
        self.fy_dropdown.addItems(fy_list)
        self.fy_dropdown.setEnabled(False)
        self.override_fy_cb.toggled.connect(self.fy_dropdown.setEnabled)
        
        # Set default selection to current FY
        current_fy = get_current_financial_year()
        if current_fy in fy_list:
            self.fy_dropdown.setCurrentIndex(fy_list.index(current_fy))
        
        fy_layout.addWidget(self.override_fy_cb)
        fy_layout.addWidget(self.fy_dropdown)
        fy_box.setLayout(fy_layout)
        
        return fy_box

    def _create_customer_section(self) -> QGroupBox:
        """Create customer and job details section."""
        customer_box = QGroupBox("Customer & Job Details")
        form1 = QFormLayout()
        
        self.customer_input = QLineEdit()
        self.customer_input.setPlaceholderText("Enter customer name")
        
        self.job_input = QLineEdit()
        self.job_input.setValidator(QIntValidator(0, Config.JOB_OP_MAX))
        self.job_input.setPlaceholderText("Enter job number")
        
        self.op_input = QLineEdit()
        self.op_input.setValidator(QIntValidator(0, Config.JOB_OP_MAX))
        self.op_input.setPlaceholderText("Enter OP number")
        
        self.product_type = QComboBox()
        self.product_type.addItems(["UPS", "Battery Charger"])
        
        form1.addRow("Customer Name:", self.customer_input)
        form1.addRow("Job Number:", self.job_input)
        form1.addRow("OP Number:", self.op_input)
        form1.addRow("Product Type:", self.product_type)
        customer_box.setLayout(form1)
        
        return customer_box

    def _create_sticker_section(self) -> QGroupBox:
        """Create sticker image selection section."""
        sticker_box = QGroupBox("Sticker Image")
        hbox = QHBoxLayout()
        
        self.sticker_path = QLineEdit()
        self.sticker_path.setPlaceholderText("Select sticker image...")
        
        browse_btn = QPushButton("Browse...")
        browse_btn.clicked.connect(self.browse_sticker)
        
        hbox.addWidget(self.sticker_path)
        hbox.addWidget(browse_btn)
        sticker_box.setLayout(hbox)
        
        return sticker_box

    def _create_ups_section(self) -> QGroupBox:
        """Create UPS configuration section."""
        ups_box = QGroupBox("UPS Configuration")
        ups_form = QFormLayout()
        
        self.num_sets = QSpinBox()
        self.num_sets.setRange(Config.UPS_SETS_MIN, Config.UPS_SETS_MAX)
        self.num_sets.setValue(1)
        
        self.ups_per_set = QSpinBox()
        self.ups_per_set.setRange(Config.UPS_PER_SET_MIN, Config.UPS_PER_SET_MAX)
        self.ups_per_set.setValue(1)
        
        self.kva_rating = QLineEdit()
        self.kva_rating.setValidator(QIntValidator(Config.KVA_MIN, Config.KVA_MAX))
        self.kva_rating.setPlaceholderText("e.g. 30 (do not add kVA)")
        
        ups_form.addRow("Number of Sets:", self.num_sets)
        ups_form.addRow("UPS per Set:", self.ups_per_set)
        ups_form.addRow("Power Rating (kVA):", self.kva_rating)
        ups_box.setLayout(ups_form)
        
        return ups_box

    def _create_charger_section(self) -> QGroupBox:
        """Create battery charger configuration section."""
        charger_box = QGroupBox("Battery Charger Configuration")
        ch_form = QFormLayout()
        
        # Show product details checkbox
        self.show_prod_label_cb = QCheckBox("Show Product Details in Sticker")
        self.show_prod_label_cb.setChecked(True)
        self.show_prod_label_cb.toggled.connect(self.update_charger_field_state)
        ch_form.addRow(self.show_prod_label_cb)
        
        # Input fields
        self.voltage = QLineEdit()
        self.voltage.setValidator(QIntValidator(1, 1000))
        self.voltage.setPlaceholderText("e.g. 48")
        
        self.current = QLineEdit()
        self.current.setValidator(QIntValidator(1, 500))
        self.current.setPlaceholderText("e.g. 10")
        
        self.battery_capacity = QLineEdit()
        self.battery_capacity.setValidator(QIntValidator(1, 5000))
        self.battery_capacity.setPlaceholderText("e.g. 100")
        
        self.charger_type = QComboBox()
        self.charger_type.addItems(["FC", "FC & FCB", "FCBC", "DFCBC"])
        
        self.battery_type = QComboBox()
        self.battery_type.addItems(["VRLA", "NICAD", "Planté", "Tubular", "Li-Ion", "Li-Po"])
        
        self.num_chargers = QSpinBox()
        self.num_chargers.setRange(Config.CHARGERS_MIN, Config.CHARGERS_MAX)
        self.num_chargers.setValue(1)
        
        ch_form.addRow("Charger Voltage (V):", self.voltage)
        ch_form.addRow("Charger Current (A):", self.current)
        ch_form.addRow("Battery Capacity (Ah):", self.battery_capacity)
        ch_form.addRow("Charger Type:", self.charger_type)
        ch_form.addRow("Battery Type:", self.battery_type)
        ch_form.addRow("Number of Chargers:", self.num_chargers)
        charger_box.setLayout(ch_form)
        
        return charger_box

    def _create_options_section(self) -> QGroupBox:
        """Create post-creation options section."""
        options_box = QGroupBox("Post Creation Options")
        opt_layout = QVBoxLayout()

        self.auto_open_cb = QCheckBox("Auto-open file after creation")
        self.auto_open_cb.setChecked(False)

        self.auto_print_cb = QCheckBox("Auto-print file after creation")
        self.auto_print_cb.setChecked(True)

        opt_layout.addWidget(self.auto_open_cb)
        opt_layout.addWidget(self.auto_print_cb)
        options_box.setLayout(opt_layout)
        
        return options_box

    def _create_generate_button(self) -> QPushButton:
        """Create the generate button."""
        generate_btn = QPushButton("Generate DOCX")
        generate_btn.clicked.connect(self.generate_docx_threaded)
        generate_btn.setStyleSheet("""
            QPushButton {
                background-color: #2f80ed;
                color: white;
                font-weight: bold;
                padding: 10px;
                border-radius: 5px;
                font-size: 11pt;
            }
            QPushButton:hover {
                background-color: #1e6cd4;
            }
            QPushButton:pressed {
                background-color: #1558b0;
            }
        """)
        
        return generate_btn

    # ---------- Event Handlers ----------
    
    def toggle_start_index(self, checked: bool) -> None:
        """Toggle BCH numbering start index."""
        self.start_index = 0 if checked else 1
        logger.info(f"BCH numbering start index set to: {self.start_index}")

    def show_about(self) -> None:
        """Display about dialog with version information."""
        # Fetch latest release info
        latest_release = get_latest_github_release()
        
        # Build version info text
        current_version_text = f"<b>{Config.VERSION}</b>"
        
        if latest_release:
            latest_version = latest_release['version']
            release_url = latest_release['url']
            
            try:
                # Compare versions
                if version.parse(Config.VERSION) < version.parse(latest_version):
                    version_status = (
                        f"<span style='color:#ff6b6b;'>⚠️ Update Available</span><br>"
                        f"Latest: <a href='{release_url}' style='color:#2F80ED; text-decoration:none;'>"
                        f"<b>v{latest_version}</b></a>"
                    )
                else:
                    version_status = "<span style='color:#51cf66;'>✓ Up to date</span>"
            except Exception as e:
                logger.warning(f"Failed to compare versions: {e}")
                version_status = (
                    f"Latest: <a href='{release_url}' style='color:#2F80ED; text-decoration:none;'>"
                    f"<b>v{latest_version}</b></a>"
                )
        else:
            version_status = "<span style='color:#999;'>Unable to check for updates</span>"
        
        about_text = (
            "<div style='font-family:Segoe UI; font-size:10pt; color:#333;'>"
            f"<h2 style='color:#2F80ED; margin-bottom:4px;'>{Config.APP_NAME}</h2>"
            f"<p style='margin:2px 0;'>Current Version: {current_version_text}</p>"
            f"<p style='margin:2px 0 8px 0;'>{version_status}</p>"
            "<hr style='border:none; border-top:1px solid #ccc; margin:8px 0;'>"
            f"<p style='margin:4px 0;'>Developed by <b>{Config.ORGANIZATION} Technologies</b></p>"
            "<p style='margin:4px 0;'>Author: <b>Amit Kumar Nandi</b></p>"
            "<p style='margin:6px 0;'>"
            "For updates, documentation, and releases, visit:<br>"
            "<a href='https://bitmutex.com' style='color:#2F80ED; text-decoration:none;'>"
            "https://bitmutex.com</a>"
            "</p>"
            "<hr style='border:none; border-top:1px solid #ccc; margin:8px 0;'>"
            "<p style='font-size:9pt; color:#777;'>© 2025 Bitmutex Technologies. All rights reserved.</p>"
            "</div>"
        )

        msg_box = QMessageBox(self)
        msg_box.setWindowTitle(f"About – {Config.APP_NAME}")
        msg_box.setTextFormat(Qt.TextFormat.RichText)
        msg_box.setIcon(QMessageBox.Icon.Information)
        msg_box.setText(about_text)
        msg_box.setStandardButtons(QMessageBox.StandardButton.Ok)
        
        # Enable links to be clickable
        msg_box.setTextInteractionFlags(
            Qt.TextInteractionFlag.TextBrowserInteraction
        )
        
        msg_box.exec()

    def update_charger_field_state(self) -> None:
        """Show/hide charger detail fields based on checkbox state."""
        show = self.show_prod_label_cb.isChecked()
        ch_form = self.show_prod_label_cb.parent().layout()
        
        # Fields that toggle visibility
        toggle_fields = [
            self.voltage,
            self.current,
            self.battery_capacity,
            self.charger_type,
            self.battery_type,
        ]
        
        # Update visibility for each field and its label
        for widget in toggle_fields:
            row = ch_form.getWidgetPosition(widget)[0]
            if row != -1:
                label_item = ch_form.itemAt(row, QFormLayout.ItemRole.LabelRole)
                field_item = ch_form.itemAt(row, QFormLayout.ItemRole.FieldRole)
                
                if label_item and label_item.widget():
                    label_item.widget().setVisible(show)
                if field_item and field_item.widget():
                    field_item.widget().setVisible(show)
        
        # Number of chargers always visible
        row = ch_form.getWidgetPosition(self.num_chargers)[0]
        if row != -1:
            label_item = ch_form.itemAt(row, QFormLayout.ItemRole.LabelRole)
            field_item = ch_form.itemAt(row, QFormLayout.ItemRole.FieldRole)
            
            if label_item and label_item.widget():
                label_item.widget().setVisible(True)
            if field_item and field_item.widget():
                field_item.widget().setVisible(True)
        
        logger.debug(f"Charger field state updated: show_details={show}")

    def update_visibility(self) -> None:
        """Update visibility of UPS/Charger sections based on product type."""
        is_ups = self.product_type.currentText() == "UPS"
        
        for box in self.findChildren(QGroupBox):
            if box.title() == "UPS Configuration":
                box.setVisible(is_ups)
            elif box.title() == "Battery Charger Configuration":
                box.setVisible(not is_ups)
        
        logger.debug(f"Product type changed to: {self.product_type.currentText()}")

    def browse_sticker(self) -> None:
        """Open file dialog to select sticker image."""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Sticker Image",
            "",
            "Image Files (*.png *.jpg *.jpeg *.bmp);;All Files (*)"
        )
        
        if file_path:
            self.sticker_path.setText(file_path)
            logger.info(f"Sticker image selected: {file_path}")

    def open_output_path(self) -> None:
        """Open the sticker output folder in file explorer."""
        try:
            # Create directory if it doesn't exist
            Config.DOCS_DIR.mkdir(parents=True, exist_ok=True)
            
            if sys.platform.startswith("win"):
                os.startfile(Config.DOCS_DIR)
            elif sys.platform == "darwin":
                subprocess.run(["open", str(Config.DOCS_DIR)])
            else:
                subprocess.run(["xdg-open", str(Config.DOCS_DIR)])
            
            logger.info(f"Opened output folder: {Config.DOCS_DIR}")
            
        except Exception as e:
            logger.error(f"Failed to open output folder: {e}")
            QMessageBox.warning(
                self,
                "Error",
                f"Could not open output folder:\n{e}"
            )

    def save_output_path(self, filename: str) -> Path:
        """
        Get full path for saving output file.
        
        Args:
            filename: Output filename
            
        Returns:
            Full path to output file
        """
        try:
            Config.DOCS_DIR.mkdir(parents=True, exist_ok=True)
            logger.debug(f"Output directory ensured: {Config.DOCS_DIR}")
        except Exception as e:
            logger.error(f"Failed to create output directory: {e}")
            # Fallback to user's Documents folder
            fallback_dir = Path.home() / "Documents"
            logger.warning(f"Using fallback directory: {fallback_dir}")
            return fallback_dir / filename
        
        return Config.DOCS_DIR / filename

    def open_github_release(self) -> None:
        """Open GitHub releases page for updates."""
        try:
            webbrowser.open("https://github.com/aamitn/sticker-generator/releases")
            logger.info("Opened GitHub releases page")
        except Exception as e:
            logger.error(f"Failed to open GitHub releases: {e}")
            QMessageBox.warning(
                self,
                "Error",
                f"Could not open GitHub releases page:\n{e}"
            )

    def purge_all_docx(self) -> None:
        """Delete all .docx files from the output folder."""
        reply = QMessageBox.question(
            self,
            "Confirm Delete",
            f"Are you sure you want to delete ALL .docx files in:\n{Config.DOCS_DIR}?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            deleted_count = 0
            failed_files = []
            
            for file in Config.DOCS_DIR.glob("*.docx"):
                try:
                    file.unlink()
                    deleted_count += 1
                    logger.info(f"Deleted: {file.name}")
                except Exception as e:
                    logger.error(f"Failed to delete {file.name}: {e}")
                    failed_files.append(file.name)

            # Show result message
            if failed_files:
                QMessageBox.warning(
                    self,
                    "Partial Success",
                    f"Deleted {deleted_count} file(s).\n\n"
                    f"Failed to delete: {', '.join(failed_files)}"
                )
            else:
                QMessageBox.information(
                    self,
                    "Deleted",
                    f"Total .docx files deleted: {deleted_count}"
                    if deleted_count else "No .docx files found."
                )

    def validate_inputs(self) -> Tuple[bool, str]:
        """
        Validate all input fields.
        
        Returns:
            Tuple of (is_valid, error_message)
        """
        product_type = self.product_type.currentText()
        sticker_path = self.sticker_path.text().strip()
        customer_name = self.customer_input.text().strip()
        job_no = self.job_input.text().strip()
        op_no = self.op_input.text().strip()

        # Check required fields
        if not all([sticker_path, customer_name, job_no, op_no]):
            return False, "Please fill all required fields and select a sticker image."

        # Check sticker file exists
        if not os.path.exists(sticker_path):
            return False, f"Sticker image not found: {sticker_path}"

        # Product-specific validation
        if product_type == "UPS":
            if not self.kva_rating.text().strip():
                return False, "Please enter KVA rating for UPS."
            
            try:
                kva = int(self.kva_rating.text().strip())
                if not (Config.KVA_MIN <= kva <= Config.KVA_MAX):
                    return False, f"KVA rating must be between {Config.KVA_MIN} and {Config.KVA_MAX}"
            except ValueError:
                return False, "Invalid KVA rating. Please enter a valid number."
        
        else:  # Battery Charger
            if self.show_prod_label_cb.isChecked():
                required_fields = {
                    "Voltage": self.voltage.text().strip(),
                    "Current": self.current.text().strip(),
                    "Battery Capacity": self.battery_capacity.text().strip(),
                }
                
                missing = [name for name, value in required_fields.items() if not value]
                if missing:
                    return False, f"Please fill in: {', '.join(missing)}"

        return True, ""

    def generate_docx_threaded(self) -> None:
        """Start DOCX generation in background thread."""
        try:
            # Validate inputs
            is_valid, error_msg = self.validate_inputs()
            if not is_valid:
                QMessageBox.warning(self, "Validation Error", error_msg)
                return

            product_type = self.product_type.currentText()
            sticker_path = self.sticker_path.text().strip()
            customer_name = self.customer_input.text().strip()
            job_no = self.job_input.text().strip()
            op_no = self.op_input.text().strip()

            # Prepare generation parameters
            kwargs = dict(
                product_type=product_type,
                sticker_path=sticker_path,
                customer_name=customer_name,
                job_no=job_no,
                op_no=op_no,
                start_index=self.start_index,
            )

            if product_type == "UPS":
                kwargs.update(
                    kva_rating=int(self.kva_rating.text().strip()),
                    num_sets=self.num_sets.value(),
                    ups_per_set=self.ups_per_set.value()
                )
            else:
                kwargs.update(
                    voltage=self.voltage.text().strip(),
                    current=self.current.text().strip(),
                    battery_capacity=self.battery_capacity.text().strip(),
                    charger_type=self.charger_type.currentText(),
                    battery_type=self.battery_type.currentText(),
                    num_chargers=self.num_chargers.value()
                )

            # Show progress dialog
            self.progress_dialog = QProgressDialog(
                "Generating DOCX...", 
                "Cancel", 
                0, 
                100, 
                self
            )
            self.progress_dialog.setWindowTitle("Please Wait")
            self.progress_dialog.setWindowModality(Qt.WindowModality.ApplicationModal)
            self.progress_dialog.setMinimumDuration(0)
            self.progress_dialog.show()

            # Start worker thread
            self.worker = DocxWorker(self, **kwargs)
            self.worker.progress.connect(self.progress_dialog.setValue)
            self.worker.finished.connect(self.on_generation_finished)
            self.worker.error.connect(self.on_generation_error)
            self.worker.start()
            
            logger.info("Started DOCX generation worker thread")

        except Exception as e:
            logger.error(f"Failed to start generation: {e}", exc_info=True)
            QMessageBox.critical(self, "Error", f"Failed to start generation:\n{e}")

    def on_generation_finished(self, output: str) -> None:
        """Handle successful DOCX generation."""
        self.progress_dialog.close()
        
        logger.info(f"Document generated successfully: {output}")
        
        QMessageBox.information(
            self,
            "Success",
            f"Document generated successfully:\n{output}"
        )
        
        # Auto-open file if enabled
        if self.auto_open_cb.isChecked():
            try:
                if sys.platform.startswith("win"):
                    os.startfile(output)
                elif sys.platform == "darwin":
                    subprocess.run(["open", output])
                else:
                    subprocess.run(["xdg-open", output])
                logger.info("Auto-opened generated document")
            except Exception as e:
                logger.error(f"Failed to auto-open document: {e}")
        
        # Auto-print if enabled
        if self.auto_print_cb.isChecked():
            self.handle_auto_print(output)

    def on_generation_error(self, error_msg: str) -> None:
        """Handle generation error."""
        self.progress_dialog.close()
        logger.error(f"Generation error: {error_msg}")
        
        QMessageBox.critical(
            self,
            "Generation Error",
            f"Failed to generate document:\n{error_msg}"
        )

    def handle_auto_print(self, output_path: str) -> None:
        """Handle automatic printing of generated document."""
        try:
            use_default_printer = self.use_default_printer_action.isChecked()
            current_platform = platform.system().lower()

            if use_default_printer:
                # Auto print with default printer
                if current_platform == "windows":
                    os.startfile(output_path, "print")
                elif current_platform in ("linux", "darwin"):
                    subprocess.run(["lp", output_path], check=True)
                else:
                    QMessageBox.information(
                        self,
                        "Info",
                        f"Automatic printing not supported on {current_platform.title()} yet."
                    )
                logger.info("Document sent to default printer")
            else:
                # Show print dialog
                self.print_docx_via_dialog(output_path)
                
        except Exception as e:
            logger.error(f"Print error: {e}")
            QMessageBox.warning(
                self,
                "Print Error",
                f"Could not print document:\n{e}"
            )

    def print_docx_via_dialog(self, docx_path: str) -> None:
        """Show print dialog and print document."""
        try:
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
            
            dialog = QPrintDialog(printer, self)
            dialog.setWindowTitle("Select Printer to Print Sticker")

            if dialog.exec():
                # User confirmed - print using system application
                current_platform = sys.platform
                
                if current_platform.startswith("win"):
                    os.startfile(docx_path, "print")
                elif current_platform == "darwin":
                    subprocess.run(["open", "-a", "Preview", docx_path])
                else:
                    subprocess.run(["xdg-open", docx_path])
                
                logger.info("Document sent to selected printer")

        except Exception as e:
            logger.error(f"Print dialog error: {e}")
            QMessageBox.warning(
                self,
                "Print Error",
                f"Could not print document:\n{e}"
            )

    # ---------- Settings Persistence ----------
    
    def load_settings(self) -> None:
        """Restore checkbox states and preferences from saved settings."""
        try:
            self.auto_open_cb.setChecked(
                self.settings.value("auto_open", False, bool)
            )
            self.auto_print_cb.setChecked(
                self.settings.value("auto_print", True, bool)
            )
            self.use_default_printer_action.setChecked(
                self.settings.value("use_default_printer", True, bool)
            )
            self.show_prod_label_cb.setChecked(
                self.settings.value("show_prod_label", True, bool)
            )
            self.start_0_action.setChecked(
                self.settings.value("start_from_zero", False, bool)
            )
            
            # Update start index based on loaded setting
            self.start_index = 0 if self.start_0_action.isChecked() else 1
            
            # Update UI based on loaded settings
            self.update_charger_field_state()
            
            logger.info("Settings loaded successfully")
            
        except Exception as e:
            logger.warning(f"Failed to load settings: {e}")

    def save_settings(self) -> None:
        """Save checkbox states and preferences."""
        try:
            self.settings.setValue("auto_open", self.auto_open_cb.isChecked())
            self.settings.setValue("auto_print", self.auto_print_cb.isChecked())
            self.settings.setValue(
                "use_default_printer",
                self.use_default_printer_action.isChecked()
            )
            self.settings.setValue(
                "show_prod_label",
                self.show_prod_label_cb.isChecked()
            )
            self.settings.setValue(
                "start_from_zero",
                self.start_0_action.isChecked()
            )
            
            logger.info("Settings saved successfully")
            
        except Exception as e:
            logger.warning(f"Failed to save settings: {e}")

    def closeEvent(self, event) -> None:
        """Handle application close event."""
        logger.info("Application closing")
        self.save_settings()
        event.accept()

    # ---------- Helper Methods ----------
    
    def get_financial_year_from_year(self, year: int) -> str:
        """Wrapper for module-level function."""
        return get_financial_year_from_year(year)

    def get_current_financial_year(self) -> str:
        """Wrapper for module-level function."""
        return get_current_financial_year()


# ----------------------------------------
# Application Entry Point
# ----------------------------------------
def main():
    """Main application entry point."""
    try:
        logger.info("=" * 60)
        logger.info(f"Starting {Config.APP_NAME} v{Config.VERSION}")
        logger.info(f"Python version: {sys.version}")
        logger.info(f"Platform: {platform.platform()}")
        logger.info(f"Executable: {sys.executable}")
        logger.info(f"Working directory: {os.getcwd()}")
        logger.info(f"App data directory: {Config.APP_DATA_DIR}")
        logger.info(f"Output directory: {Config.DOCS_DIR}")
        logger.info("=" * 60)
        
        # Ensure critical directories exist
        try:
            Config.APP_DATA_DIR.mkdir(parents=True, exist_ok=True)
            Config.DOCS_DIR.mkdir(parents=True, exist_ok=True)
            logger.info("Application directories initialized successfully")
        except Exception as e:
            logger.warning(f"Could not create all directories: {e}")
        
        app = QApplication(sys.argv)
        app.setApplicationName(Config.APP_NAME)
        app.setOrganizationName(Config.ORGANIZATION)
        app.setApplicationVersion(Config.VERSION)
        
        window = StickerApp()
        window.show()
        
        exit_code = app.exec()
        logger.info(f"Application exited with code: {exit_code}")
        sys.exit(exit_code)
        
    except PermissionError as e:
        error_msg = (
            f"Permission Error: {e}\n\n"
            f"The application doesn't have permission to access required directories.\n"
            f"Please try running the application with appropriate permissions."
        )
        logger.critical(error_msg, exc_info=True)
        
        # Show error dialog if possible
        try:
            app = QApplication.instance() or QApplication(sys.argv)
            QMessageBox.critical(None, "Permission Error", error_msg)
        except:
            print(error_msg, file=sys.stderr)
        
        sys.exit(1)
        
    except Exception as e:
        error_msg = f"Fatal error: {e}"
        logger.critical(error_msg, exc_info=True)
        
        # Show error dialog if possible
        try:
            app = QApplication.instance() or QApplication(sys.argv)
            QMessageBox.critical(
                None,
                "Fatal Error",
                f"An unexpected error occurred:\n\n{e}\n\n"
                f"Please check the log file at:\n{Config.LOG_FILE}"
            )
        except:
            print(error_msg, file=sys.stderr)
        
        sys.exit(1)


if __name__ == "__main__":
    main()